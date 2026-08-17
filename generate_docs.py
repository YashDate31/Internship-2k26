import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
from reportlab.lib.units import inch

# Target directories
BASE_DIR = r"c:\Users\Yash\OneDrive\Desktop\Internship 2026\Final Project\resourses 5th sem"
ASSIGNMENT_DIR = os.path.join(BASE_DIR, "assignment")
IMP_DIR = os.path.join(BASE_DIR, "imp")

os.makedirs(ASSIGNMENT_DIR, exist_ok=True)
os.makedirs(IMP_DIR, exist_ok=True)

# Helper function to set table cell background color in docx
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_docx_document(title, subject, subject_code, doc_type, content_data, output_path):
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title Banner Table
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "1E3A8A") # Navy Blue
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("MAHARASHTRA STATE BOARD OF TECHNICAL EDUCATION")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(217, 119, 6) # Amber gold

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(f"DEPARTMENT OF COMPUTER ENGINEERING — SEMESTER 5")
    run2.font.name = "Arial"
    run2.font.size = Pt(10)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(255, 255, 255)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(12)
    run3 = p3.add_run(f"{subject.upper()} ({subject_code}) — {doc_type.upper()}")
    run3.font.name = "Arial"
    run3.font.size = Pt(14)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Metadata Block
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    meta_data = [
        [f"Subject: {subject} ({subject_code})", "Semester: 5th Sem (CO5I)"],
        [f"Document: {doc_type}", "Academic Year: 2025 - 2026"]
    ]
    
    for r_idx, row in enumerate(meta_table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = Inches(3.4)
            set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(meta_data[r_idx][c_idx])
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Sections Content
    for sec in content_data:
        # Heading
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        run_h = h.add_run(sec["section_title"])
        run_h.font.name = "Arial"
        run_h.font.size = Pt(12.5)
        run_h.font.bold = True
        run_h.font.color.rgb = RGBColor(30, 58, 138)

        # Questions & Answers
        for qa in sec["items"]:
            q_p = doc.add_paragraph()
            q_p.paragraph_format.space_before = Pt(8)
            q_p.paragraph_format.space_after = Pt(3)
            run_q = q_p.add_run(f"Q: {qa['question']}")
            run_q.font.name = "Arial"
            run_q.font.size = Pt(10.5)
            run_q.font.bold = True
            run_q.font.color.rgb = RGBColor(15, 23, 42)

            if "marks" in qa:
                run_m = q_p.add_run(f"  [{qa['marks']}]")
                run_m.font.name = "Arial"
                run_m.font.size = Pt(9.5)
                run_m.font.bold = True
                run_m.font.color.rgb = RGBColor(217, 119, 6)

            a_p = doc.add_paragraph()
            a_p.paragraph_format.space_before = Pt(2)
            a_p.paragraph_format.space_after = Pt(8)
            a_p.paragraph_format.left_indent = Inches(0.2)
            run_a = a_p.add_run(f"Ans: {qa['answer']}")
            run_a.font.name = "Calibri"
            run_a.font.size = Pt(10)
            run_a.font.color.rgb = RGBColor(51, 65, 85)

    doc.save(output_path)
    print(f"Created DOCX: {output_path}")

def create_pdf_document(title, subject, subject_code, doc_type, content_data, output_path):
    pdf = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.6*inch,
        leftMargin=0.6*inch,
        topMargin=0.6*inch,
        bottomMargin=0.6*inch
    )

    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'HeaderTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=14,
        textColor=colors.HexColor('#FFFFFF'),
        alignment=1,
        spaceAfter=4
    )
    
    subtitle_style = ParagraphStyle(
        'HeaderSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        textColor=colors.HexColor('#D97706'),
        alignment=1,
        spaceAfter=4
    )

    sec_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=12,
        spaceAfter=6
    )

    q_style = ParagraphStyle(
        'QuestionText',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=6,
        spaceAfter=2
    )

    a_style = ParagraphStyle(
        'AnswerText',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        textColor=colors.HexColor('#334155'),
        leftIndent=12,
        spaceBefore=2,
        spaceAfter=6
    )

    story = []

    # Header Box Table
    header_data = [
        [Paragraph('MAHARASHTRA STATE BOARD OF TECHNICAL EDUCATION', subtitle_style)],
        [Paragraph(f'DEPARTMENT OF COMPUTER ENGINEERING — SEMESTER 5', title_style)],
        [Paragraph(f'{subject.upper()} ({subject_code}) — {doc_type.upper()}', title_style)]
    ]
    
    header_table = Table(header_data, colWidths=[7.0*inch])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    
    story.append(header_table)
    story.append(Spacer(1, 0.15*inch))

    # Meta Table
    meta_content = [
        [Paragraph(f'<b>Subject:</b> {subject} ({subject_code})', styles['Normal']), Paragraph('<b>Semester:</b> 5th Semester (CO5I)', styles['Normal'])],
        [Paragraph(f'<b>Document Type:</b> {doc_type}', styles['Normal']), Paragraph('<b>Academic Year:</b> 2025 - 2026', styles['Normal'])]
    ]
    meta_table = Table(meta_content, colWidths=[3.5*inch, 3.5*inch])
    meta_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F8FAFC')),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#E2E8F0')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.2*inch))

    # Content
    for sec in content_data:
        story.append(Paragraph(sec['section_title'], sec_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#1E3A8A'), spaceAfter=8))
        
        for qa in sec['items']:
            marks_str = f"  <font color='#D97706'>[{qa['marks']}]</font>" if "marks" in qa else ""
            story.append(Paragraph(f"Q: {qa['question']}{marks_str}", q_style))
            story.append(Paragraph(f"<b>Ans:</b> {qa['answer']}", a_style))

    pdf.build(story)
    print(f"Created PDF: {output_path}")

print("Document generator helper script initialized.")
